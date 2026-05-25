"""
Generate AegisFX_Status_Report.docx with embedded screenshots.

Usage:
    1. Save your screenshots to docs/screenshots/ with the filenames listed
       in the SCREENSHOTS dict below (e.g. 01_pnl_panel.png).
    2. Run: python docs/generate_report.py
    3. Open: docs/AegisFX_Status_Report.docx

Any screenshot that is missing will be skipped with a "[Image missing]" note.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

THIS_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(THIS_DIR, "screenshots")
OUTPUT_PATH = os.path.join(THIS_DIR, "AegisFX_Status_Report.docx")

# Maps section -> (caption, screenshot filename)
SCREENSHOTS = {
    "pnl": ("Figure 1 — P&L Panel and Performance KPIs", "01_pnl_panel.png"),
    "daily": ("Figure 2 — Daily Performance Table", "02_daily_performance.png"),
    "equity": ("Figure 3 — Equity Curve", "03_equity_curve.png"),
    "positions": ("Figure 4 — Current Positions and Risk Exposure", "04_positions_risk.png"),
    "ai_agreement": ("Figure 5 — AI Agreement Panel", "05_ai_agreement.png"),
    "proposals": ("Figure 6 — AI Trade Proposals", "06_ai_proposals.png"),
    "queue": ("Figure 7 — AI Approval Queue", "07_approval_queue.png"),
    "analytics": ("Figure 8 — AI Proposal Analytics", "08_proposal_analytics.png"),
    "attribution": ("Figure 9 — Strategy Attribution by Regime", "09_strategy_attribution.png"),
    "accuracy": ("Figure 10 — AI Recommendation Accuracy", "10_recommendation_accuracy.png"),
    "exports": ("Figure 11 — Investor Report Exports", "11_investor_exports.png"),
    "trend": ("Figure 12 — AI Confidence Trend", "12_confidence_trend.png"),
    "alerts": ("Figure 13 — Alerts and System Status", "13_alerts_status.png"),
}


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_image(doc, key):
    """Embed screenshot if present; otherwise note that it's missing."""
    caption, filename = SCREENSHOTS[key]
    path = os.path.join(SCREENSHOTS_DIR, filename)

    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
    else:
        warn = doc.add_paragraph()
        warn_run = warn.add_run(f"[Image missing: {filename}]")
        warn_run.italic = True
        warn_run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, rows):
    """rows is a list of lists; first row is the header."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True


def build_report():
    doc = Document()

    # --- Title page ---
    title = doc.add_heading("AegisFX Trading System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Status Report")
    sub_run.bold = True
    sub_run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Date: 2026-05-25\n").italic = True
    meta.add_run("Environment: OANDA Demo Account\n").italic = True
    meta.add_run("Author: Sunday Dada (Colaberry Internship)").italic = True

    doc.add_page_break()

    # --- Executive Summary ---
    add_heading(doc, "Executive Summary", level=1)
    add_paragraph(doc,
        "The AegisFX automated trading system is live and operational, connected to "
        "OANDA's demo trading account. It executes trades through a deterministic "
        "governance pipeline with AI-assisted market analysis powered by OpenAI. "
        "The dashboard provides real-time portfolio state, AI analysis, performance "
        "analytics, and full operator control."
    )
    add_paragraph(doc,
        "This document walks through the current state of the system, panel by panel, "
        "as observed on the live dashboard."
    )

    # --- 1. P&L ---
    add_heading(doc, "1. Portfolio Health (P&L Panel)", level=1)
    add_image(doc, "pnl")

    add_table(doc, [
        ["Metric", "Value", "Interpretation"],
        ["Unrealized P&L", "-$0.03", "Currently open positions are slightly negative"],
        ["Action Signal", "LOSS (yellow)", "Open positions unprofitable, no drawdown warning"],
        ["Account Balance", "$99,999.91", "Demo account (started at $100,000)"],
        ["Drawdown from Peak", "0.00%", "No drawdown from equity peak"],
        ["Equity", "$99,999.87", "Balance + Unrealized P&L"],
    ])

    add_paragraph(doc, "Status: Account essentially flat. Demo capital fully intact.", bold=True)

    # --- 2. KPIs ---
    add_heading(doc, "2. Performance KPIs", level=1)
    add_paragraph(doc,
        "The KPI row directly below the P&L panel summarizes lifetime performance "
        "from closed trades."
    )

    add_table(doc, [
        ["Metric", "Value", "Interpretation"],
        ["Closed Trades", "37", "37 completed trade cycles"],
        ["Win Rate", "43.2%", "16 of 37 trades were profitable"],
        ["Total Profit", "-$312.69", "Cumulative realized loss across all closed trades"],
        ["Total Pips", "-312.7", "Net pip movement against the system"],
    ])

    add_paragraph(doc,
        "Honest assessment: Most activity to date came from random test trades "
        "generated by dry_run_sustained.py, not AI-curated trades. This is the "
        "baseline before AI-driven proposals start contributing real signal."
    )

    # --- 3. Daily Performance ---
    add_heading(doc, "3. Daily Performance Table", level=1)
    add_image(doc, "daily")

    add_table(doc, [
        ["Date", "Trades", "Win %", "Pips", "Profit"],
        ["2026-04-11", "12", "41.7%", "+314.69", "+314.69"],
        ["2026-04-14", "3", "33.3%", "-1.34", "-1.34"],
        ["2026-04-15", "14", "71.4%", "+6.86", "+6.86"],
        ["2026-05-04", "4", "0.0%", "-316.53", "-316.53"],
        ["2026-05-25", "4", "0.0%", "-316.37", "-316.37"],
    ])

    add_paragraph(doc,
        "Key insight: The system had strong days (Apr 11, Apr 15) offset by two bad "
        "sessions (May 4, May 25). The bad days were random-trade test sessions during "
        "unfavorable conditions. This consistency view is what investors examine to "
        "gauge stability over time."
    )

    # --- 4. Equity Curve ---
    add_heading(doc, "4. Equity Curve", level=1)
    add_image(doc, "equity")

    add_paragraph(doc, "The cumulative realized profit chart shows:")
    for bullet in [
        "Rapid climb to ~$315 from April 11 winning session",
        "Held there for approximately 3 weeks",
        "Sharp drop in May from two losing sessions",
        "Net ending near zero",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    add_paragraph(doc,
        "Interpretation: System is not yet alpha-positive — gains and losses cancel. "
        "This is the baseline state before AI proposal trading begins.", italic=True
    )

    # --- 5. Current Positions ---
    add_heading(doc, "5. Current Positions and Risk Exposure", level=1)
    add_image(doc, "positions")

    add_paragraph(doc, "Two open positions held on OANDA demo account:", bold=True)
    add_table(doc, [
        ["Request ID", "Pair", "Direction", "Units", "Entry Price"],
        ["SUSTAINED-36768b40-...", "GBP/USD", "Long", "1", "1.3567"],
        ["SUSTAINED-f225e9db-...", "GBP/USD", "Long", "1", "1.3561"],
    ])

    add_paragraph(doc, "Risk Exposure metrics:", bold=True)
    add_table(doc, [
        ["Metric", "Value"],
        ["Signal", "LOW (green)"],
        ["Total Exposure", "4.0"],
        ["Max Allowed", "100.0"],
        ["Utilization", "4.0%"],
    ])

    add_paragraph(doc, "Net Exposure by Currency:", bold=True)
    add_table(doc, [
        ["Currency", "Exposure"],
        ["GBP", "+2"],
        ["USD", "-2"],
    ])

    add_paragraph(doc,
        "Status: System is using only 4% of risk budget. Plenty of headroom available. "
        "The maximum was recently raised from 10 to 100 to allow larger position sizing "
        "if needed."
    )

    # --- 7. AI Agreement ---
    add_heading(doc, "6. AI Agreement (Core AI Signal)", level=1)
    add_image(doc, "ai_agreement")

    add_paragraph(doc,
        "This is the most important AI panel. It shows OpenAI's real-time market "
        "interpretation combined with the deterministic strategy recommendation layer."
    )

    add_table(doc, [
        ["Field", "Value"],
        ["Signal", "STRONG (green)"],
        ["Current Regime", "Ranging"],
        ["Model Confidence", "85%"],
        ["Risk Mode", "NORMAL (green)"],
        ["Execution", "ALLOWED (green)"],
        ["Recommended Strategy", "MeanReversion_v1"],
        ["Trade Bias", "NEUTRAL"],
    ])

    add_paragraph(doc, "AI Summary (from OpenAI):", bold=True)
    add_paragraph(doc,
        '"Currency pairs show flat trends with low to medium volatility indicating a ranging market."',
        italic=True
    )

    add_paragraph(doc, "Per-Pair Analysis (from OpenAI):", bold=True)
    for line in [
        "EUR/USD — Flat trend and low volatility suggest limited directional movement",
        "GBP/USD — Flat trend with medium volatility indicates sideways price action with some fluctuations",
        "USD/JPY — Flat trend and low volatility confirm stable price levels without strong momentum",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_paragraph(doc,
        "What's Happening: OpenAI receives real intraday candle data (sourced from Alpha "
        "Vantage), classifies the market as Ranging with high confidence (85%), and the "
        "deterministic strategy layer recommends MeanReversion_v1. However, the Trade Bias "
        "is NEUTRAL because mean reversion requires price to be diverging from the mean in "
        "a direction — and currently the market is too flat."
    )
    add_paragraph(doc,
        'This is correct, conservative behavior. The AI is honestly saying "no edge available '
        'right now" instead of forcing trades.', bold=True
    )

    # --- 8. AI Proposals ---
    add_heading(doc, "7. AI Trade Proposals", level=1)
    add_image(doc, "proposals")
    add_paragraph(doc,
        'Status: "No AI trade proposals available." Why empty: The strategy recommendation '
        "produced NEUTRAL bias, so the proposal generator correctly outputs zero proposals. "
        "The system refuses to generate trade ideas when there is no signal. This is "
        "governance working as designed."
    )

    # --- 9. Approval Queue ---
    add_heading(doc, "8. AI Approval Queue", level=1)
    add_image(doc, "queue")
    add_paragraph(doc,
        'Status: "No proposals pending approval." This panel will populate when the AI '
        "detects a tradable regime (e.g., Trending), strategy bias becomes directional "
        "(LONG or SHORT), and proposals get queued for operator decision."
    )

    # --- 10. Proposal Analytics ---
    add_heading(doc, "9. AI Proposal Analytics", level=1)
    add_image(doc, "analytics")
    add_paragraph(doc,
        "All metrics currently at zero — no AI proposals have been executed yet."
    )
    add_paragraph(doc, "Approval Funnel: Proposed: 0 | Approved: 0 | Executed: 0", bold=True)
    add_paragraph(doc,
        "This panel becomes meaningful once the operator starts approving and executing "
        "AI suggestions."
    )

    # --- 11. Strategy Attribution ---
    add_heading(doc, "10. Strategy Attribution by Regime", level=1)
    add_image(doc, "attribution")
    add_paragraph(doc,
        'Status: "No attribution data yet — needs executed AI trades that have closed." '
        "Will show per-(regime, strategy) performance once AI trades close. The purpose "
        "of this panel is to reveal which strategy performs best in which market regime — "
        "a critical signal for refining or retiring strategies over time."
    )

    # --- 12. Accuracy ---
    add_heading(doc, "11. AI Recommendation Accuracy", level=1)
    add_image(doc, "accuracy")
    add_paragraph(doc,
        'Status: "No executed AI recommendations have closed yet." This panel will measure '
        "how often executed AI recommendations turn out profitable — the ultimate validation "
        "of whether the AI adds value."
    )

    # --- 13. Exports ---
    add_heading(doc, "12. Investor Report Exports", level=1)
    add_image(doc, "exports")
    add_paragraph(doc, "Two operational download buttons:", bold=True)
    for line in [
        "Download Investor Report (CSV) — Consolidated performance data",
        "Download Investor Report (HTML) — Browser-printable PDF-style investor report",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc,
        "Both export consolidated metrics suitable for sharing with stakeholders. The HTML "
        "report can be printed to PDF directly from any browser."
    )

    # --- 14. Confidence Trend ---
    add_heading(doc, "13. AI Confidence Trend", level=1)
    add_image(doc, "trend")
    add_paragraph(doc,
        "A line chart showing AI confidence over time. Currently displays a flat line at "
        '~85%, indicating the AI has been consistently classifying "Ranging" with high '
        "confidence throughout this session."
    )
    add_paragraph(doc,
        "Recent Regime Changes & Summaries table shows 5 entries, all from 2026-05-25 "
        'with identical "Ranging at 85%" entries — confirming the market has been stably '
        "ranging for the recent observation window."
    )

    # --- 15. Alerts ---
    add_heading(doc, "14. Alerts and System Status", level=1)
    add_image(doc, "alerts")

    add_table(doc, [
        ["Indicator", "Value", "Color"],
        ["Overall", "ALL CLEAR", "Green"],
        ["Trading Enabled", "ON (operator toggle)", "Green"],
        ["Circuit Breaker", "INACTIVE", "Green"],
        ["Broker Connection", "CONNECTED", "Green"],
        ["Pending Trades", "0", "Green"],
        ["Last Successful Trade", "2026-05-04 08:04:52", "—"],
        ["Active Alerts", "None", "—"],
    ])

    add_paragraph(doc,
        "Status: All system gates are green. Broker connected, no pending or stuck trades, "
        "operator has trading enabled, no active alerts.", bold=True
    )

    # --- Architecture ---
    add_heading(doc, "Architecture Summary", level=1)
    add_paragraph(doc, "AegisFX is built on four governance principles:", bold=True)
    for i, principle in enumerate([
        "AI advises; deterministic code decides. OpenAI provides regime analysis. "
        "Deterministic Python rules translate that analysis into strategy recommendations "
        "and trade proposals.",

        "Two human clicks separate AI from execution. Operator must explicitly Approve a "
        "proposal, then explicitly Execute it. No autonomous trade flow exists.",

        "Every action passes deterministic gates. Even approved AI proposals must pass risk "
        "evaluation, position netting, broker health check, rate limiting, and circuit "
        "breaker before reaching the broker.",

        "State persists and reconciles. SQLite-backed ledger survives restarts. On startup, "
        "the system queries the broker for any PENDING trades and resolves them.",
    ], start=1):
        doc.add_paragraph(f"{i}. {principle}", style="List Number")

    # --- What This Means ---
    add_heading(doc, "What This Means", level=1)

    add_heading(doc, "The System Is Working", level=2)
    for line in [
        "All deterministic gates function correctly",
        "AI is producing real, consistent signal from live market data (Alpha Vantage → OpenAI)",
        "AI correctly abstains in flat markets rather than forcing trades",
        "Broker integration is solid (OANDA demo connected, 2 real positions held)",
        "Operator controls (close buttons, toggles, approval queue) all functional",
        "Audit trail is intact (integrity check passes — 0 outstanding pending trades)",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "What's Pending", level=2)
    for line in [
        "No AI-executed trades have closed yet. Analytics panels for proposal analytics, "
        "strategy attribution, and recommendation accuracy need real AI execution data to populate.",
        "Current P&L reflects random-trade testing, not AI-curated trading. The AI trading "
        "layer is what's expected to add edge.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "Recommended Next Steps", level=2)
    for i, line in enumerate([
        "Continue running the dashboard during active forex market hours",
        "When the AI detects a Trending or Volatile regime with directional bias, real "
        "proposals will appear in the queue",
        "Approve high-confidence proposals, reject low-confidence ones, and execute "
        "approved ones",
        "After a series of AI trades close, the analytics panels will reveal whether the AI "
        "is adding positive edge versus baseline",
    ], start=1):
        doc.add_paragraph(f"{i}. {line}", style="List Number")

    add_heading(doc, "For Management", level=2)
    add_paragraph(doc,
        "The system is in a state ready for live demo evaluation. The full pipeline — "
        "market data ingestion, AI regime analysis, deterministic risk governance, broker "
        "execution, and operator dashboard — works end-to-end. Real-money deployment would "
        "require additional gates per the deployment checklist (extended dry-run validation, "
        "broker reconciliation testing, etc.)."
    )

    # --- Appendix ---
    add_heading(doc, "Appendix — Technical Stack", level=1)
    add_table(doc, [
        ["Layer", "Technology"],
        ["Market Data", "Alpha Vantage REST API (5-minute intraday candles)"],
        ["AI Analysis", "OpenAI (gpt-4.1-mini)"],
        ["Broker", "OANDA REST API (practice account)"],
        ["State Storage", "SQLite with WAL journaling"],
        ["Dashboard", "Streamlit (Python)"],
        ["Orchestration", "Pure Python (deterministic governance pipeline)"],
    ])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Report generated from live dashboard state on 2026-05-25.")
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
    print(f"Report saved to: {OUTPUT_PATH}")

    # Report which screenshots are missing
    missing = []
    for key, (caption, filename) in SCREENSHOTS.items():
        path = os.path.join(SCREENSHOTS_DIR, filename)
        if not os.path.exists(path):
            missing.append(filename)

    if missing:
        print(f"\nMissing screenshots (placeholders inserted in document):")
        for m in missing:
            print(f"  - docs/screenshots/{m}")
        print(f"\nSave your screenshots with these filenames in:")
        print(f"  {SCREENSHOTS_DIR}")
        print(f"Then re-run this script to regenerate with embedded images.")
    else:
        print("\nAll 13 screenshots found and embedded.")
